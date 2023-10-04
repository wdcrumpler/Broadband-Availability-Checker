import time
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document

data = pd.read_csv('addressesXfinity.csv')
data['Eligible (Xfinity)'] = ''
data['Has Account (Xfinity)'] = ''
data['Need Unit Number (Xfinity)'] = ''
data['No Service (Xfinity)'] = ''
data['Unknown (Xfinity)']=''
data['Traceback']=''

for index, row in data.iterrows():
    address = row['Address']
    zipcode = row['ZipCode']
    fullAddress = row['Full Address']

    try:

# Load Spectrum website
        # Set up Selenium Chrome driver
        chrome_options = Options()
        #chrome_options.add_argument('--headless')
        service = Service('C:\Program Files\Chrome Driver\chromedriver.exe')
        driver = webdriver.Chrome(service=service, options=chrome_options)
        driver.get('https://www.xfinity.com/overview')

#Find and fill the address and zip fields, then submit the form
        print("Full address is",fullAddress)
        time.sleep(2)

        fullAddressInputArray = WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.XPATH, '//input[@class="input text contained sc-prism-input-default" and @name="localizationAddressField"]')))
        full_address_field = fullAddressInputArray[0]
        driver.execute_script("arguments[0].value = arguments[1];", full_address_field, fullAddress)

        time.sleep(2)

        submitButtonArray = WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.XPATH, "//button[contains(., 'Check availability')]")))
        submit_button =  submitButtonArray[0]
        driver.execute_script("arguments[0].click();", submit_button)
        print ("button pressed")

        time.sleep(3)

        wait = WebDriverWait(driver, 10)
        wait.until(EC.presence_of_element_located((By.TAG_NAME, 'body')))
        content = driver.page_source 
        url = driver.current_url

        if '<span class="localization-container__header" data-testid="localization-fallback-final-header">Hmm, that address wasn' in content:
            print ("Address not found first time, re-submitting...")
            addressInputArray = WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.XPATH, '//input[@class="input text contained sc-prism-input-default" and @name="localizationAddressField"]')))
            address_field = addressInputArray[0]
            driver.execute_script("arguments[0].value = arguments[1];", address_field, address)
            
            zipInputArray = WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.XPATH, '//input[@class="input text contained sc-prism-input-default" and @name="localizationZipField"]')))
            zip_field = zipInputArray[0]
            driver.execute_script("arguments[0].value = arguments[1];", zip_field, zipcode)
            
            content = driver.page_source 
            document = Document()
            document.add_paragraph(content)
            document.save("page_source.docx")

            time.sleep(10)
            submitButtonArray = WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.XPATH, "//button[contains(., 'Check availability')]")))
            submit_button =  submitButtonArray[0]
            driver.execute_script("arguments[0].click();", submit_button)
            print ("Re-submitted...")
            time.sleep(3)
            url = driver.current_url
            content = driver.page_source 
        
        if "Is this the correct address?" in content:
            content = driver.page_source 
            document = Document()
            document.add_paragraph(content)
            document.save("page_source.docx")
            submitButtonArray = WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.XPATH, "//button[contains(., 'Yes, check availability')]")))
            submit_button =  submitButtonArray[0]
            driver.execute_script("arguments[0].click();", submit_button)
            time.sleep(3)
            url = driver.current_url

        if "Is this a home or business address?" in content:
            submitButtonArray = WebDriverWait(driver, 10).until(EC.presence_of_all_elements_located((By.XPATH, "//button[contains(., 'Home')]")))
            submit_button =  submitButtonArray[0]
            driver.execute_script("arguments[0].click();", submit_button)
            time.sleep(3)
            url = driver.current_url

        if "active-address" in url:
            data.at[index, 'Has Account (Xfinity)'] = 'X'
            print("Has Account set to X for",address)
        elif "out-of-footprint" in url:
            data.at[index, 'No Service (Xfinity)'] = 'X'
            print("No Service set to X for",address)
        elif "offers/plan-builder" in url:
            data.at[index, 'Eligible (Xfinity)'] = 'X'
            print("Eligible set to X for",address)
        #elif "mdu-fallback-container__button-container" in content:
        elif '<h2 class="localization-container__header">Is one of these the correct address?</h2>' in content:
            data.at[index, 'Need Unit Number (Xfinity)'] = 'X'
            print("Need Unit Number set to X for",address)
        else:
            data.at[index, 'Unknown (Spectrum)'] = 'X'
            print("Unknown set to X for",address)
        
        time.sleep(3)

        data.iloc[[index]].to_csv('addresses_with_info_Xfinity.csv', index=False, header=False, mode='a')

        driver.quit()

    except Exception as e:
        print(f"Exception occurred for address: {address}")
        print(str(e))
        data.at[index, 'Traceback'] = 'X'

    finally:
        driver.quit()

print("CSV file updated successfully.")